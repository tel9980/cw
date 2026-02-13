"""
Collection Letter Generator Module

Generates professional collection letters for overdue receivables in Word document format.
Supports three template types:
- FIRST_REMINDER: Polite first reminder (30 days overdue)
- SECOND_NOTICE: Firmer second notice (60 days overdue)
- FINAL_WARNING: Final warning before legal action (90 days overdue)
"""

import logging
from datetime import date, datetime
from decimal import Decimal
from enum import Enum
from typing import Optional, Dict, Any
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ..models.core_models import Counterparty


# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class LetterTemplate(Enum):
    """催款函模板类型"""
    FIRST_REMINDER = "first_reminder"  # 首次提醒（30天）
    SECOND_NOTICE = "second_notice"    # 二次催收（60天）
    FINAL_WARNING = "final_warning"    # 最后通知（90天）


class CollectionLetterGenerator:
    """
    催款函生成器
    
    生成专业的催款函Word文档，支持三种模板：
    - 首次提醒：礼貌的第一次提醒
    - 二次催收：更坚定的第二次通知
    - 最后通知：法律行动前的最后警告
    """
    
    def __init__(
        self,
        company_name: str = "贵公司",
        company_address: str = "",
        company_phone: str = "",
        company_contact: str = "",
        output_dir: str = "collection_letters"
    ):
        """
        初始化催款函生成器
        
        Args:
            company_name: 公司名称
            company_address: 公司地址
            company_phone: 公司电话
            company_contact: 联系人
            output_dir: 输出目录
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx库未安装，无法生成催款函。请运行: pip install python-docx"
            )
        
        self.company_name = company_name
        self.company_address = company_address
        self.company_phone = company_phone
        self.company_contact = company_contact
        self.output_dir = Path(output_dir)
        
        # Create output directory if it doesn't exist
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_collection_letter(
        self,
        customer: Counterparty,
        overdue_days: int,
        amount: Decimal,
        invoice_numbers: Optional[list] = None,
        due_date: Optional[date] = None,
        template_type: Optional[LetterTemplate] = None
    ) -> str:
        """
        生成催款函
        
        Args:
            customer: 客户信息
            overdue_days: 逾期天数
            amount: 欠款金额
            invoice_numbers: 发票号列表（可选）
            due_date: 原到期日（可选）
            template_type: 模板类型（可选，自动根据逾期天数选择）
        
        Returns:
            str: 生成的Word文档路径
        """
        # Auto-select template based on overdue days if not specified
        if template_type is None:
            template_type = self._select_template(overdue_days)
        
        logger.info(
            f"生成催款函: 客户={customer.name}, "
            f"逾期={overdue_days}天, 金额={amount}, 模板={template_type.value}"
        )
        
        # Create document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # Add header with company info
        self._add_header(doc)
        
        # Add title
        self._add_title(doc, template_type)
        
        # Add recipient info
        self._add_recipient_info(doc, customer)
        
        # Add letter body based on template
        self._add_letter_body(
            doc, 
            template_type, 
            customer, 
            overdue_days, 
            amount,
            invoice_numbers,
            due_date
        )
        
        # Add payment details table
        self._add_payment_details(doc, amount, invoice_numbers, due_date)
        
        # Add closing
        self._add_closing(doc, template_type)
        
        # Add signature
        self._add_signature(doc)
        
        # Add footer
        self._add_footer(doc)
        
        # Save document
        filename = self._generate_filename(customer, template_type)
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        
        logger.info(f"催款函已生成: {filepath}")
        return str(filepath)
    
    def _select_template(self, overdue_days: int) -> LetterTemplate:
        """根据逾期天数自动选择模板"""
        if overdue_days >= 90:
            return LetterTemplate.FINAL_WARNING
        elif overdue_days >= 60:
            return LetterTemplate.SECOND_NOTICE
        else:
            return LetterTemplate.FIRST_REMINDER
    
    def _add_header(self, doc: Document) -> None:
        """添加文档头部（公司信息）"""
        # Company name
        p = doc.add_paragraph()
        run = p.add_run(self.company_name)
        run.font.size = Pt(14)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Company details
        if self.company_address or self.company_phone:
            p = doc.add_paragraph()
            details = []
            if self.company_address:
                details.append(f"地址: {self.company_address}")
            if self.company_phone:
                details.append(f"电话: {self.company_phone}")
            run = p.add_run(" | ".join(details))
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_title(self, doc: Document, template_type: LetterTemplate) -> None:
        """添加标题"""
        titles = {
            LetterTemplate.FIRST_REMINDER: "付款提醒函",
            LetterTemplate.SECOND_NOTICE: "催款通知函",
            LetterTemplate.FINAL_WARNING: "最后催款通知函"
        }
        
        p = doc.add_paragraph()
        run = p.add_run(titles[template_type])
        run.font.size = Pt(16)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
    
    def _add_recipient_info(self, doc: Document, customer: Counterparty) -> None:
        """添加收件人信息"""
        # Date
        p = doc.add_paragraph()
        run = p.add_run(f"日期: {date.today().strftime('%Y年%m月%d日')}")
        run.font.size = Pt(10)
        
        # Recipient
        p = doc.add_paragraph()
        run = p.add_run(f"致: {customer.name}")
        run.font.size = Pt(11)
        run.font.bold = True
        
        if customer.contact_person:
            p = doc.add_paragraph()
            run = p.add_run(f"收件人: {customer.contact_person}")
            run.font.size = Pt(10)
        
        doc.add_paragraph()
    
    def _add_letter_body(
        self,
        doc: Document,
        template_type: LetterTemplate,
        customer: Counterparty,
        overdue_days: int,
        amount: Decimal,
        invoice_numbers: Optional[list],
        due_date: Optional[date]
    ) -> None:
        """添加信函正文"""
        # Greeting
        p = doc.add_paragraph()
        run = p.add_run(f"尊敬的{customer.name}:")
        run.font.size = Pt(11)
        
        doc.add_paragraph()
        
        # Body content based on template
        if template_type == LetterTemplate.FIRST_REMINDER:
            self._add_first_reminder_body(doc, customer, overdue_days, amount, due_date)
        elif template_type == LetterTemplate.SECOND_NOTICE:
            self._add_second_notice_body(doc, customer, overdue_days, amount, due_date)
        else:  # FINAL_WARNING
            self._add_final_warning_body(doc, customer, overdue_days, amount, due_date)
        
        doc.add_paragraph()
    
    def _add_first_reminder_body(
        self,
        doc: Document,
        customer: Counterparty,
        overdue_days: int,
        amount: Decimal,
        due_date: Optional[date]
    ) -> None:
        """添加首次提醒正文"""
        paragraphs = [
            f"感谢贵司一直以来对{self.company_name}的支持与信任。",
            "",
            f"根据我司记录，贵司有一笔金额为 {amount:,.2f} 元的应付款项已逾期 {overdue_days} 天"
            + (f"（原到期日: {due_date.strftime('%Y年%m月%d日')}）" if due_date else "")
            + "。",
            "",
            "我们理解在业务运营中可能会出现各种情况，如果贵司已经安排付款，请忽略此函。"
            "如果尚未安排，我们恳请贵司尽快处理此笔款项。",
            "",
            "如有任何疑问或需要讨论付款安排，请随时与我们联系。我们期待继续与贵司保持良好的合作关系。"
        ]
        
        for text in paragraphs:
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.line_spacing = 1.5
                for run in p.runs:
                    run.font.size = Pt(11)
            else:
                doc.add_paragraph()
    
    def _add_second_notice_body(
        self,
        doc: Document,
        customer: Counterparty,
        overdue_days: int,
        amount: Decimal,
        due_date: Optional[date]
    ) -> None:
        """添加二次催收正文"""
        paragraphs = [
            f"我司此前已就贵司的逾期款项发送过提醒函，但截至目前尚未收到付款或回复。",
            "",
            f"贵司目前有一笔金额为 {amount:,.2f} 元的应付款项已逾期 {overdue_days} 天"
            + (f"（原到期日: {due_date.strftime('%Y年%m月%d日')}）" if due_date else "")
            + "。",
            "",
            "长期的逾期付款不仅影响我司的正常运营，也可能影响双方的合作关系。"
            "我们诚恳地要求贵司在收到本函后的 7 个工作日内完成付款。",
            "",
            "如果贵司在付款方面遇到困难，请立即与我们联系，我们愿意讨论合理的付款安排方案。"
            "但如果继续未能收到付款或回复，我们将不得不采取进一步措施。"
        ]
        
        for text in paragraphs:
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.line_spacing = 1.5
                for run in p.runs:
                    run.font.size = Pt(11)
            else:
                doc.add_paragraph()
    
    def _add_final_warning_body(
        self,
        doc: Document,
        customer: Counterparty,
        overdue_days: int,
        amount: Decimal,
        due_date: Optional[date]
    ) -> None:
        """添加最后通知正文"""
        paragraphs = [
            f"这是我司就贵司逾期款项发出的最后通知函。",
            "",
            f"贵司有一笔金额为 {amount:,.2f} 元的应付款项已严重逾期 {overdue_days} 天"
            + (f"（原到期日: {due_date.strftime('%Y年%m月%d日')}）" if due_date else "")
            + "。尽管我司多次发函提醒，但至今未收到任何付款或正式回复。",
            "",
            "鉴于上述情况，我司郑重要求贵司在收到本函后的 3 个工作日内全额支付该笔款项。"
            "如果在此期限内仍未收到付款，我司将不得不采取以下措施：",
            "",
            "1. 暂停与贵司的所有业务往来；",
            "2. 将此事提交法律部门，通过法律途径追讨欠款；",
            "3. 将贵司的信用记录报告给相关信用机构。",
            "",
            "我们不希望事态发展到这一步，恳请贵司立即处理此事。如需讨论付款安排，"
            "请在 24 小时内与我司联系。"
        ]
        
        for text in paragraphs:
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.line_spacing = 1.5
                for run in p.runs:
                    run.font.size = Pt(11)
                    # Make warning text bold
                    if "最后通知" in text or "法律途径" in text or "3 个工作日" in text:
                        run.font.bold = True
            else:
                doc.add_paragraph()
    
    def _add_payment_details(
        self,
        doc: Document,
        amount: Decimal,
        invoice_numbers: Optional[list],
        due_date: Optional[date]
    ) -> None:
        """添加付款明细表"""
        # Add section title
        p = doc.add_paragraph()
        run = p.add_run("欠款明细:")
        run.font.size = Pt(11)
        run.font.bold = True
        
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "项目"
        header_cells[1].text = "内容"
        header_cells[2].text = "金额（元）"
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        
        # Add invoice numbers if provided
        if invoice_numbers:
            row_cells = table.add_row().cells
            row_cells[0].text = "发票号"
            row_cells[1].text = ", ".join(invoice_numbers)
            row_cells[2].text = ""
        
        # Add due date if provided
        if due_date:
            row_cells = table.add_row().cells
            row_cells[0].text = "原到期日"
            row_cells[1].text = due_date.strftime('%Y年%m月%d日')
            row_cells[2].text = ""
        
        # Add total amount
        row_cells = table.add_row().cells
        row_cells[0].text = "欠款总额"
        row_cells[1].text = ""
        row_cells[2].text = f"{amount:,.2f}"
        
        # Make total row bold
        for paragraph in row_cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        for paragraph in row_cells[2].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(192, 0, 0)  # Red color
        
        doc.add_paragraph()
    
    def _add_closing(self, doc: Document, template_type: LetterTemplate) -> None:
        """添加结束语"""
        if template_type == LetterTemplate.FIRST_REMINDER:
            closing = "再次感谢贵司的理解与配合，期待早日收到贵司的付款。"
        elif template_type == LetterTemplate.SECOND_NOTICE:
            closing = "感谢贵司的重视与配合，期待在规定期限内收到贵司的付款。"
        else:  # FINAL_WARNING
            closing = "希望贵司能够重视此事，避免不必要的法律纠纷。"
        
        p = doc.add_paragraph(closing)
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.size = Pt(11)
        
        doc.add_paragraph()
    
    def _add_signature(self, doc: Document) -> None:
        """添加签名"""
        p = doc.add_paragraph()
        run = p.add_run("此致")
        run.font.size = Pt(11)
        
        p = doc.add_paragraph()
        run = p.add_run("敬礼！")
        run.font.size = Pt(11)
        
        doc.add_paragraph()
        
        # Company signature
        p = doc.add_paragraph()
        run = p.add_run(self.company_name)
        run.font.size = Pt(11)
        run.font.bold = True
        
        if self.company_contact:
            p = doc.add_paragraph()
            run = p.add_run(f"联系人: {self.company_contact}")
            run.font.size = Pt(10)
        
        if self.company_phone:
            p = doc.add_paragraph()
            run = p.add_run(f"联系电话: {self.company_phone}")
            run.font.size = Pt(10)
        
        p = doc.add_paragraph()
        run = p.add_run(f"日期: {date.today().strftime('%Y年%m月%d日')}")
        run.font.size = Pt(10)
    
    def _add_footer(self, doc: Document) -> None:
        """添加页脚"""
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        run = p.add_run("---")
        run.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = doc.add_paragraph()
        run = p.add_run("本函由系统自动生成，如有疑问请联系财务部门")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _generate_filename(
        self,
        customer: Counterparty,
        template_type: LetterTemplate
    ) -> str:
        """生成文件名"""
        template_names = {
            LetterTemplate.FIRST_REMINDER: "首次提醒",
            LetterTemplate.SECOND_NOTICE: "二次催收",
            LetterTemplate.FINAL_WARNING: "最后通知"
        }
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        template_name = template_names[template_type]
        
        # Sanitize customer name for filename
        safe_name = "".join(c for c in customer.name if c.isalnum() or c in (' ', '-', '_'))
        safe_name = safe_name.strip()[:50]  # Limit length
        
        return f"催款函_{safe_name}_{template_name}_{timestamp}.docx"
