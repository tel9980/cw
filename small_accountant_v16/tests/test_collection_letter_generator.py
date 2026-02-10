"""
Unit tests for CollectionLetterGenerator

Tests the generation of collection letters in Word document format
with three different templates.
"""

import pytest
from datetime import date, datetime, timedelta
from decimal import Decimal
from pathlib import Path
import tempfile
import shutil

from small_accountant_v16.models.core_models import Counterparty, CounterpartyType
from small_accountant_v16.reminders.collection_letter_generator import (
    CollectionLetterGenerator,
    LetterTemplate
)

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@pytest.fixture
def temp_output_dir():
    """Create a temporary directory for test outputs"""
    temp_dir = tempfile.mkdtemp()
    yield temp_dir
    # Cleanup
    shutil.rmtree(temp_dir, ignore_errors=True)


@pytest.fixture
def generator(temp_output_dir):
    """Create a CollectionLetterGenerator instance for testing"""
    return CollectionLetterGenerator(
        company_name="测试科技有限公司",
        company_address="北京市朝阳区测试路123号",
        company_phone="010-12345678",
        company_contact="张经理",
        output_dir=temp_output_dir
    )


@pytest.fixture
def sample_customer():
    """Create a sample customer for testing"""
    return Counterparty(
        id="CUST001",
        name="优质客户有限公司",
        type=CounterpartyType.CUSTOMER,
        contact_person="李总",
        phone="021-87654321",
        email="lizong@example.com",
        address="上海市浦东新区客户路456号",
        tax_id="91310000123456789X",
        created_at=datetime.now(),
        updated_at=datetime.now()
    )


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
class TestCollectionLetterGenerator:
    """Test suite for CollectionLetterGenerator"""
    
    def test_initialization(self, temp_output_dir):
        """Test generator initialization"""
        generator = CollectionLetterGenerator(
            company_name="测试公司",
            company_address="测试地址",
            company_phone="123456",
            company_contact="测试联系人",
            output_dir=temp_output_dir
        )
        
        assert generator.company_name == "测试公司"
        assert generator.company_address == "测试地址"
        assert generator.company_phone == "123456"
        assert generator.company_contact == "测试联系人"
        assert generator.output_dir == Path(temp_output_dir)
        assert generator.output_dir.exists()
    
    def test_output_directory_creation(self, temp_output_dir):
        """Test that output directory is created if it doesn't exist"""
        non_existent_dir = Path(temp_output_dir) / "new_dir"
        assert not non_existent_dir.exists()
        
        generator = CollectionLetterGenerator(output_dir=str(non_existent_dir))
        assert non_existent_dir.exists()
    
    def test_template_selection_30_days(self, generator):
        """Test automatic template selection for 30 days overdue"""
        template = generator._select_template(30)
        assert template == LetterTemplate.FIRST_REMINDER
    
    def test_template_selection_60_days(self, generator):
        """Test automatic template selection for 60 days overdue"""
        template = generator._select_template(60)
        assert template == LetterTemplate.SECOND_NOTICE
    
    def test_template_selection_90_days(self, generator):
        """Test automatic template selection for 90 days overdue"""
        template = generator._select_template(90)
        assert template == LetterTemplate.FINAL_WARNING
    
    def test_template_selection_edge_cases(self, generator):
        """Test template selection at boundary values"""
        assert generator._select_template(59) == LetterTemplate.FIRST_REMINDER
        assert generator._select_template(60) == LetterTemplate.SECOND_NOTICE
        assert generator._select_template(89) == LetterTemplate.SECOND_NOTICE
        assert generator._select_template(90) == LetterTemplate.FINAL_WARNING
        assert generator._select_template(120) == LetterTemplate.FINAL_WARNING
    
    def test_generate_first_reminder_letter(self, generator, sample_customer):
        """Test generating a first reminder letter"""
        filepath = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=35,
            amount=Decimal("50000.00"),
            invoice_numbers=["INV-2024-001", "INV-2024-002"],
            due_date=date.today() - timedelta(days=35),
            template_type=LetterTemplate.FIRST_REMINDER
        )
        
        # Check file was created
        assert Path(filepath).exists()
        assert filepath.endswith(".docx")
        assert "首次提醒" in filepath
        assert sample_customer.name[:10] in filepath or "优质客户" in filepath
        
        # Verify document content
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        assert "付款提醒函" in text
        assert sample_customer.name in text
        assert "50,000.00" in text
        assert "35 天" in text
        assert generator.company_name in text
        
        # Check invoice numbers in table
        table_text = "\n".join([cell.text for table in doc.tables for row in table.rows for cell in row.cells])
        assert "INV-2024-001" in table_text
    
    def test_generate_second_notice_letter(self, generator, sample_customer):
        """Test generating a second notice letter"""
        filepath = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=65,
            amount=Decimal("75000.50"),
            invoice_numbers=["INV-2024-003"],
            due_date=date.today() - timedelta(days=65),
            template_type=LetterTemplate.SECOND_NOTICE
        )
        
        # Check file was created
        assert Path(filepath).exists()
        assert "二次催收" in filepath
        
        # Verify document content
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        assert "催款通知函" in text
        assert sample_customer.name in text
        assert "75,000.50" in text
        assert "65 天" in text
        assert "7 个工作日" in text
        assert "进一步措施" in text
    
    def test_generate_final_warning_letter(self, generator, sample_customer):
        """Test generating a final warning letter"""
        filepath = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=95,
            amount=Decimal("100000.00"),
            invoice_numbers=["INV-2024-004", "INV-2024-005"],
            due_date=date.today() - timedelta(days=95),
            template_type=LetterTemplate.FINAL_WARNING
        )
        
        # Check file was created
        assert Path(filepath).exists()
        assert "最后通知" in filepath
        
        # Verify document content
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        assert "最后催款通知函" in text
        assert sample_customer.name in text
        assert "100,000.00" in text
        assert "95 天" in text
        assert "3 个工作日" in text
        assert "法律途径" in text
        assert "暂停与贵司的所有业务往来" in text
    
    def test_generate_letter_without_invoice_numbers(self, generator, sample_customer):
        """Test generating a letter without invoice numbers"""
        filepath = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=40,
            amount=Decimal("25000.00"),
            invoice_numbers=None,
            due_date=date.today() - timedelta(days=40)
        )
        
        assert Path(filepath).exists()
        
        # Document should still be valid
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "25,000.00" in text
    
    def test_generate_letter_without_due_date(self, generator, sample_customer):
        """Test generating a letter without due date"""
        filepath = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=50,
            amount=Decimal("30000.00"),
            invoice_numbers=["INV-2024-006"],
            due_date=None
        )
        
        assert Path(filepath).exists()
        
        # Document should still be valid
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "30,000.00" in text
        
        # Check invoice numbers in table
        table_text = "\n".join([cell.text for table in doc.tables for row in table.rows for cell in row.cells])
        assert "INV-2024-006" in table_text
    
    def test_generate_letter_auto_template_selection(self, generator, sample_customer):
        """Test automatic template selection based on overdue days"""
        # 30 days - should use FIRST_REMINDER
        filepath1 = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=30,
            amount=Decimal("10000.00")
        )
        assert "首次提醒" in filepath1
        
        # 60 days - should use SECOND_NOTICE
        filepath2 = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=60,
            amount=Decimal("10000.00")
        )
        assert "二次催收" in filepath2
        
        # 90 days - should use FINAL_WARNING
        filepath3 = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=90,
            amount=Decimal("10000.00")
        )
        assert "最后通知" in filepath3
    
    def test_document_structure(self, generator, sample_customer):
        """Test that generated document has proper structure"""
        filepath = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=45,
            amount=Decimal("50000.00"),
            invoice_numbers=["INV-001"],
            due_date=date.today() - timedelta(days=45)
        )
        
        doc = Document(filepath)
        
        # Check that document has content
        assert len(doc.paragraphs) > 10
        
        # Check that document has tables (payment details)
        assert len(doc.tables) > 0
        
        # Check table structure
        table = doc.tables[0]
        assert len(table.rows) >= 2  # At least header + one data row
        assert len(table.columns) == 3  # 项目, 内容, 金额
    
    def test_amount_formatting(self, generator, sample_customer):
        """Test that amounts are properly formatted with commas"""
        filepath = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=30,
            amount=Decimal("1234567.89")
        )
        
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        # Check for comma-separated formatting
        assert "1,234,567.89" in text
    
    def test_customer_name_sanitization(self, generator):
        """Test that customer names with special characters are sanitized in filename"""
        customer_with_special_chars = Counterparty(
            id="CUST002",
            name="特殊/字符\\公司:名称*",
            type=CounterpartyType.CUSTOMER,
            contact_person="王总",
            phone="123456",
            email="test@example.com",
            address="测试地址",
            tax_id="123456789",
            created_at=datetime.now(),
            updated_at=datetime.now()
        )
        
        filepath = generator.generate_collection_letter(
            customer=customer_with_special_chars,
            overdue_days=30,
            amount=Decimal("10000.00")
        )
        
        # Filename should not contain special characters
        filename = Path(filepath).name
        assert "/" not in filename
        assert "\\" not in filename
        assert ":" not in filename
        assert "*" not in filename
    
    def test_multiple_letters_unique_filenames(self, generator, sample_customer):
        """Test that multiple letters generate unique filenames"""
        import time
        
        filepath1 = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=30,
            amount=Decimal("10000.00")
        )
        
        # Small delay to ensure different timestamp
        time.sleep(1)
        
        # Generate another letter
        filepath2 = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=30,
            amount=Decimal("10000.00")
        )
        
        # Filenames should be different (due to timestamp)
        assert filepath1 != filepath2
        assert Path(filepath1).exists()
        assert Path(filepath2).exists()
    
    def test_company_info_in_document(self, generator, sample_customer):
        """Test that company information appears in the document"""
        filepath = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=30,
            amount=Decimal("10000.00")
        )
        
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        assert generator.company_name in text
        assert generator.company_address in text
        assert generator.company_phone in text
        assert generator.company_contact in text
    
    def test_customer_info_in_document(self, generator, sample_customer):
        """Test that customer information appears in the document"""
        filepath = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=30,
            amount=Decimal("10000.00")
        )
        
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        assert sample_customer.name in text
        assert sample_customer.contact_person in text
    
    def test_date_formatting(self, generator, sample_customer):
        """Test that dates are properly formatted in Chinese"""
        due_date = date(2024, 3, 15)
        
        filepath = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=30,
            amount=Decimal("10000.00"),
            due_date=due_date
        )
        
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        # Check for Chinese date format
        assert "2024年03月15日" in text or "2024年3月15日" in text
    
    def test_letter_tone_progression(self, generator, sample_customer):
        """Test that letter tone becomes firmer with each template"""
        # First reminder - polite
        filepath1 = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=30,
            amount=Decimal("10000.00"),
            template_type=LetterTemplate.FIRST_REMINDER
        )
        doc1 = Document(filepath1)
        text1 = "\n".join([p.text for p in doc1.paragraphs])
        assert "感谢" in text1
        assert "恳请" in text1
        
        # Second notice - firmer
        filepath2 = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=60,
            amount=Decimal("10000.00"),
            template_type=LetterTemplate.SECOND_NOTICE
        )
        doc2 = Document(filepath2)
        text2 = "\n".join([p.text for p in doc2.paragraphs])
        assert "要求" in text2
        assert "进一步措施" in text2
        
        # Final warning - firm
        filepath3 = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=90,
            amount=Decimal("10000.00"),
            template_type=LetterTemplate.FINAL_WARNING
        )
        doc3 = Document(filepath3)
        text3 = "\n".join([p.text for p in doc3.paragraphs])
        assert "最后通知" in text3
        assert "法律途径" in text3
        assert "暂停" in text3


@pytest.mark.skipif(DOCX_AVAILABLE, reason="Test for missing dependency")
def test_missing_docx_library():
    """Test that appropriate error is raised when python-docx is not installed"""
    with pytest.raises(ImportError, match="python-docx"):
        CollectionLetterGenerator()


class TestEdgeCases:
    """Test edge cases and error handling"""
    
    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
    def test_very_large_amount(self, generator, sample_customer):
        """Test handling of very large amounts"""
        filepath = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=30,
            amount=Decimal("999999999.99")
        )
        
        assert Path(filepath).exists()
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "999,999,999.99" in text
    
    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
    def test_small_amount(self, generator, sample_customer):
        """Test handling of small amounts"""
        filepath = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=30,
            amount=Decimal("0.01")
        )
        
        assert Path(filepath).exists()
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "0.01" in text
    
    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
    def test_very_long_customer_name(self, generator):
        """Test handling of very long customer names"""
        long_name_customer = Counterparty(
            id="CUST003",
            name="这是一个非常非常非常非常非常非常非常非常非常非常长的公司名称用于测试文件名长度限制",
            type=CounterpartyType.CUSTOMER,
            contact_person="测试",
            phone="123456",
            email="test@example.com",
            address="测试地址",
            tax_id="123456789",
            created_at=datetime.now(),
            updated_at=datetime.now()
        )
        
        filepath = generator.generate_collection_letter(
            customer=long_name_customer,
            overdue_days=30,
            amount=Decimal("10000.00")
        )
        
        # Filename should be truncated but still valid
        assert Path(filepath).exists()
        filename = Path(filepath).name
        assert len(filename) < 200  # Reasonable filename length
    
    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
    def test_many_invoice_numbers(self, generator, sample_customer):
        """Test handling of many invoice numbers"""
        many_invoices = [f"INV-2024-{i:04d}" for i in range(1, 21)]
        
        filepath = generator.generate_collection_letter(
            customer=sample_customer,
            overdue_days=30,
            amount=Decimal("10000.00"),
            invoice_numbers=many_invoices
        )
        
        assert Path(filepath).exists()
        doc = Document(filepath)
        
        # Check that invoices are in the table
        table_text = "\n".join([cell.text for table in doc.tables for row in table.rows for cell in row.cells])
        assert "INV-2024-0001" in table_text
        assert "INV-2024-0020" in table_text
